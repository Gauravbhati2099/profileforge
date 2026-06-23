from docx import Document
from zipfile import ZipFile
from PIL import Image
import pytesseract
import io


def extract_headers_and_footers(doc):

    content = []

    for section in doc.sections:

        try:

            header = section.header

            for para in header.paragraphs:

                text = para.text.strip()

                if text:

                    content.append(
                        f"[HEADER] {text}"
                    )

        except Exception:
            pass

        try:

            footer = section.footer

            for para in footer.paragraphs:

                text = para.text.strip()

                if text:

                    content.append(
                        f"[FOOTER] {text}"
                    )

        except Exception:
            pass

    return content


def extract_paragraphs(doc):

    content = []

    for para in doc.paragraphs:

        text = para.text.strip()

        if text:

            content.append(text)

    return content


def extract_tables(doc):

    content = []

    for table in doc.tables:

        content.append("[TABLE START]")

        for row in table.rows:

            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
            )

            if row_text.strip():

                content.append(row_text)

        content.append("[TABLE END]")

    return content


def extract_images(file_path):

    content = []

    try:

        with ZipFile(file_path) as archive:

            for name in archive.namelist():

                if not name.startswith(
                    "word/media/"
                ):
                    continue

                try:

                    image_bytes = archive.read(name)

                    image = Image.open(
                        io.BytesIO(image_bytes)
                    )

                    text = (
                        pytesseract.image_to_string(
                            image
                        )
                    )

                    if text.strip():

                        content.append(
                            f"[IMAGE OCR: {name}]"
                        )

                        content.append(
                            text.strip()
                        )

                except Exception as e:

                    print(
                        f"[IMAGE OCR SKIP] {name}: {e}"
                    )

    except Exception as e:

        print(
            f"[IMAGE EXTRACTION ERROR] {e}"
        )

    return content


def extract_docx(file_path: str) -> str:

    content = []

    try:

        doc = Document(file_path)

        # Headers
        content.extend(
            extract_headers_and_footers(doc)
        )

        # Paragraphs
        content.extend(
            extract_paragraphs(doc)
        )

        # Tables
        content.extend(
            extract_tables(doc)
        )

    except Exception as e:

        print(
            f"[DOCX CONTENT ERROR] {file_path}: {e}"
        )

    # Images OCR
    content.extend(
        extract_images(file_path)
    )

    return "\n".join(content)